import os, json, csv, re
from datetime import datetime
from typing import Optional, Dict, Any, List

# 選用依賴（若未安裝，會自動略過或給出人性化錯誤）
def _try_import(name):
    try:
        return __import__(name)
    except Exception:
        return None

python_docx = _try_import("docx")              # 解析 .docx
docx2txt     = _try_import("docx2txt")         # 解析 .docx（替代）
textract     = _try_import("textract")         # 嘗試解析 .doc（舊格式）
pdfplumber   = _try_import("pdfplumber")       # 解析 PDF（建議）
PyPDF2       = _try_import("PyPDF2")           # PDF 備用方案

class FileToMarkdownJSON:
    """
    讀取 doc/docx/pdf/csv/txt -> 轉 Markdown -> 存 JSON/JSONL
    - convert(path): 回傳單一檔案的 dict 結構（包含 markdown）
    - convert_many(paths): 回傳多檔案 list
    - save_json(data, out_path): 輸出 JSON 或 JSONL
    """

    SUPPORTED_EXTS = {".doc", ".docx", ".pdf", ".csv", ".txt"}

    def __init__(self, page_separator: str = "\n\n---\n\n"):
        self.page_separator = page_separator

    # ---------- Public APIs ----------
    def convert(self, path: str) -> Dict[str, Any]:
        ext = os.path.splitext(path)[1].lower()
        if ext not in self.SUPPORTED_EXTS:
            raise ValueError(f"不支援的副檔名：{ext}，支援 {sorted(self.SUPPORTED_EXTS)}")

        if ext == ".txt":
            markdown = self._txt_to_md(path)
        elif ext == ".csv":
            markdown = self._csv_to_md(path)
        elif ext == ".docx":
            markdown = self._docx_to_md(path)
        elif ext == ".doc":
            markdown = self._doc_to_md(path)
        elif ext == ".pdf":
            markdown = self._pdf_to_md(path)
        else:
            raise ValueError(f"未處理的副檔名：{ext}")

        return {
            "filename": os.path.basename(path),
            "filepath": os.path.abspath(path),
            "filetype": ext.lstrip("."),
            "markdown": markdown,
            "metadata": {
                "converted_at": datetime.utcnow().isoformat() + "Z",
                "pages": self._count_pages(path, ext, markdown),
                "size_bytes": os.path.getsize(path) if os.path.exists(path) else None,
            },
        }

    def convert_many(self, paths: List[str]) -> List[Dict[str, Any]]:
        results = []
        for p in paths:
            try:
                results.append(self.convert(p))
            except Exception as e:
                results.append({
                    "filename": os.path.basename(p),
                    "filepath": os.path.abspath(p),
                    "error": str(e)
                })
        return results

    def save_json(self, data: List[Dict[str, Any]] | Dict[str, Any], out_path: str) -> None:
        """
        - 若 out_path 以 .json 結尾：輸出為單一 JSON（若 data 是 list 就存成陣列）
        - 若 out_path 以 .jsonl 結尾：每行一筆 JSON
        """
        os.makedirs(os.path.dirname(os.path.abspath(out_path)), exist_ok=True)

        if out_path.lower().endswith(".jsonl"):
            if isinstance(data, dict):
                data = [data]
            with open(out_path, "w", encoding="utf-8") as f:
                for row in data:
                    f.write(json.dumps(row, ensure_ascii=False) + "\n")
        else:
            with open(out_path, "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2)

    # ---------- Helpers: TXT ----------
    def _txt_to_md(self, path: str) -> str:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
        # 簡單清洗：將連續空白壓成最多兩行
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    # ---------- Helpers: CSV -> Markdown Table ----------
    def _csv_to_md(self, path: str) -> str:
        def escape_pipe(s: str) -> str:
            return s.replace("|", r"\|")
        rows = []
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            reader = csv.reader(f)
            for row in reader:
                rows.append([escape_pipe(col.strip()) for col in row])

        if not rows:
            return "_(此 CSV 為空)_"

        header = rows[0]
        body = rows[1:] if len(rows) > 1 else []

        md = []
        md.append("| " + " | ".join(header) + " |")
        md.append("| " + " | ".join(["---"] * len(header)) + " |")
        for r in body:
            # 若行長度不齊，補空字串
            if len(r) < len(header):
                r = r + [""] * (len(header) - len(r))
            md.append("| " + " | ".join(r) + " |")
        return "\n".join(md)

    # ---------- Helpers: DOCX ----------
    def _docx_to_md(self, path: str) -> str:
        # 優先用 python-docx 取得更結構化資訊
        if python_docx is None and docx2txt is None:
            raise ImportError("請安裝 python-docx 或 docx2txt 用於解析 .docx（pip install python-docx 或 pip install docx2txt）")

        if python_docx is not None:
            from docx.enum.text import WD_BREAK
            doc = python_docx.Document(path)

            blocks = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                # 依據粗體/樣式簡單推斷為標題（非常簡化）
                is_heading = getattr(para.style, "name", "").lower().startswith("heading")
                is_bold = any(run.bold for run in para.runs if run.text.strip())
                if is_heading or is_bold:
                    blocks.append(f"## {text}")
                else:
                    blocks.append(text)

            # 表格轉 Markdown
            for table in doc.tables:
                blocks.append(self._docx_table_to_md(table))

            md = "\n\n".join(blocks).strip()
            return md if md else "_(此 DOCX 無內容)_"

        # 後備方案：docx2txt（純文字）
        text = docx2txt.process(path) or ""
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip() if text.strip() else "_(此 DOCX 無內容)_"

    def _docx_table_to_md(self, table) -> str:
        rows = []
        for row in table.rows:
            cols = [self._inline_clean(cell.text) for cell in row.cells]
            rows.append(cols)
        # 去重複（docx 表格常有 cell 合併導致重複）
        uniq = []
        for r in rows:
            if not uniq or r != uniq[-1]:
                uniq.append(r)
        if not uniq:
            return ""
        header = uniq[0]
        body = uniq[1:]
        out = []
        out.append("| " + " | ".join(header) + " |")
        out.append("| " + " | ".join(["---"] * len(header)) + " |")
        for r in body:
            if len(r) < len(header):
                r = r + [""] * (len(header) - len(r))
            out.append("| " + " | ".join(r) + " |")
        return "\n".join(out)

    # ---------- Helpers: DOC (legacy) ----------
    def _doc_to_md(self, path: str) -> str:
        if textract is None:
            raise ImportError("解析 .doc 需要 textract 或 antiword，建議安裝：pip install textract（需系統層依賴 antiword/catdoc）")
        try:
            raw = textract.process(path)  # 可能回傳 bytes
            text = raw.decode("utf-8", errors="ignore")
        except Exception as e:
            raise RuntimeError(f".doc 解析失敗：{e}")
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip() if text.strip() else "_(此 DOC 無內容)_"

    # ---------- Helpers: PDF ----------
    def _pdf_to_md(self, path: str) -> str:
        # 優先 pdfplumber（較穩定的文字座標解析）
        if pdfplumber is not None:
            pages_md = []
            with pdfplumber.open(path) as pdf:
                for i, page in enumerate(pdf.pages, start=1):
                    text = page.extract_text() or ""
                    text = text.strip()
                    if not text:
                        continue
                    pages_md.append(f"### 第 {i} 頁\n\n{text}")
            return self.page_separator.join(pages_md).strip() if pages_md else "_(未從 PDF 解析到文字，可能是掃描影像)_"

        # 備用：PyPDF2（純文字）
        if PyPDF2 is not None:
            reader = PyPDF2.PdfReader(path)
            pages_md = []
            for i, page in enumerate(reader.pages, start=1):
                text = page.extract_text() or ""
                text = text.strip()
                if text:
                    pages_md.append(f"### 第 {i} 頁\n\n{text}")
            return self.page_separator.join(pages_md).strip() if pages_md else "_(未從 PDF 解析到文字，可能是掃描影像)_"

        raise ImportError("請安裝 pdfplumber 或 PyPDF2 以解析 PDF（pip install pdfplumber 或 pip install PyPDF2）")

    # ---------- Utilities ----------
    def _inline_clean(self, s: str) -> str:
        s = (s or "").strip()
        s = re.sub(r"\s+", " ", s)
        s = s.replace("|", r"\|")
        return s

    def _count_pages(self, path: str, ext: str, markdown: str) -> Optional[int]:
        try:
            if ext == ".pdf":
                if pdfplumber is not None:
                    with pdfplumber.open(path) as pdf:
                        return len(pdf.pages)
                if PyPDF2 is not None:
                    reader = PyPDF2.PdfReader(path)
                    return len(reader.pages)
            # 其他格式沒有天然頁數概念，回傳 None
            return None
        except Exception:
            return None


# ---------------- 示例用法 ----------------
if __name__ == "__main__":
    """
    範例：
    python convert.py
    會示範把同資料夾下的檔案轉成 JSON 與 JSONL
    """
    demo_files = [
        "sample.doc", "sample.docx", "sample.pdf", "sample.csv", "sample.txt"
    ]
    demo_files = [f for f in demo_files if os.path.exists(f)]

    converter = FileToMarkdownJSON()

    if demo_files:
        results = converter.convert_many(demo_files)
        # 輸出單一 JSON（陣列）
        converter.save_json(results, "./output/converted.json")
        # 輸出 JSONL（每行一筆）
        converter.save_json(results, "./output/converted.jsonl")
        print("轉檔完成：./output/converted.json 與 ./output/converted.jsonl")
    else:
        print("請將 sample.doc/docx/pdf/csv/txt 放在同目錄後再執行示例。")
